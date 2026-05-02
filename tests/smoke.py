#!/usr/bin/env python3
"""
Tests de fumée Orion — vérifie tout ce qui ne demande pas de commande vocale.

Usage :
    python tests/smoke.py                 # tous les tests
    python tests/smoke.py --no-network    # skip web/serveur (offline)
    python tests/smoke.py --image-gen     # inclut generate_image (coûte 1 appel Gemini)
    python tests/smoke.py --server-only   # uniquement les tests d'endpoints serveur

Prérequis : .env présent (le script le lit automatiquement).
Pour les tests d'endpoints : `python start.py server` doit tourner en parallèle.

Sortie : pass / skip / fail par test, récap final, exit code = nb d'échecs.
"""
from __future__ import annotations

import argparse
import json
import os
import sys
import tempfile
import time
import urllib.error
import urllib.request
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

# Force UTF-8 sur stdout/stderr (Windows cp1252 ne sait pas afficher Unicode)
for _stream_name in ("stdout", "stderr"):
    _stream = getattr(sys, _stream_name, None)
    if hasattr(_stream, "reconfigure"):
        try:
            _stream.reconfigure(encoding="utf-8", errors="replace")
        except Exception:
            pass

# ─── Charge .env ────────────────────────────────────────────────────────
try:
    from dotenv import load_dotenv
    load_dotenv(ROOT / ".env")
except ImportError:
    if (ROOT / ".env").exists():
        for line in (ROOT / ".env").read_text(encoding="utf-8").splitlines():
            line = line.strip()
            if line and not line.startswith("#") and "=" in line:
                k, v = line.split("=", 1)
                os.environ.setdefault(k.strip(), v.strip().strip('"').strip("'"))

from branding import get_env, sync_env_aliases  # noqa: E402

sync_env_aliases()


# ─── Couleurs ──────────────────────────────────────────────────────────
if os.name == "nt":
    os.system("")  # active ANSI sous Windows 10+
GREEN, RED, YELLOW, GRAY, CYAN, RESET = (
    "\033[32m", "\033[31m", "\033[33m", "\033[90m", "\033[36m", "\033[0m"
)


# ─── Runner ────────────────────────────────────────────────────────────
class Runner:
    def __init__(self):
        self.passed = 0
        self.failed = 0
        self.skipped = 0
        self.failures: list[tuple[str, str]] = []
        self.start = time.time()

    def section(self, title: str):
        print(f"\n{CYAN}━━━ {title} ━━━{RESET}")

    def test(self, name: str, fn):
        t0 = time.time()
        try:
            result = fn()
        except SkipTest as exc:
            elapsed = (time.time() - t0) * 1000
            print(f"  {YELLOW}SKIP{RESET}  {name:<55s} {GRAY}{elapsed:>5.0f}ms — {exc}{RESET}")
            self.skipped += 1
            return
        except AssertionError as exc:
            elapsed = (time.time() - t0) * 1000
            msg = str(exc) or "assertion failed"
            print(f"  {RED}FAIL{RESET}  {name:<55s} {GRAY}{elapsed:>5.0f}ms — {msg}{RESET}")
            self.failed += 1
            self.failures.append((name, msg))
            return
        except Exception as exc:
            elapsed = (time.time() - t0) * 1000
            msg = f"{type(exc).__name__}: {exc}"
            print(f"  {RED}FAIL{RESET}  {name:<55s} {GRAY}{elapsed:>5.0f}ms — {msg}{RESET}")
            self.failed += 1
            self.failures.append((name, msg))
            return
        elapsed = (time.time() - t0) * 1000
        info = f" — {result}" if isinstance(result, str) and result else ""
        print(f"  {GREEN}PASS{RESET}  {name:<55s} {GRAY}{elapsed:>5.0f}ms{info}{RESET}")
        self.passed += 1

    def report(self) -> int:
        elapsed = time.time() - self.start
        total = self.passed + self.failed + self.skipped
        print(f"\n{CYAN}━━━ Récap ━━━{RESET}")
        print(f"  Total      : {total} tests en {elapsed:.1f}s")
        print(f"  {GREEN}Passés     : {self.passed}{RESET}")
        print(f"  {YELLOW}Skip       : {self.skipped}{RESET}")
        print(f"  {RED}Échoués    : {self.failed}{RESET}")
        if self.failures:
            print(f"\n{RED}Détail des échecs :{RESET}")
            for name, msg in self.failures:
                print(f"  • {name}\n    {GRAY}{msg}{RESET}")
        return self.failed


class SkipTest(Exception):
    pass


# ═══════════════════════════════════════════════════════════════════════
# TESTS
# ═══════════════════════════════════════════════════════════════════════

def test_imports_and_registry(r: Runner):
    r.section("Imports & registry")

    def import_tools():
        from server.tools import ALL_HANDLERS
        assert len(ALL_HANDLERS) >= 30, f"Seulement {len(ALL_HANDLERS)} tools (attendu >= 30)"
        return f"{len(ALL_HANDLERS)} tools enregistrés"
    r.test("import server.tools (ALL_HANDLERS)", import_tools)

    def import_orchestrator():
        from server.orchestrator import TOOLS, _build_system_prompt
        assert len(TOOLS) >= 30
        prompt_default = _build_system_prompt(None)
        prompt_voice = _build_system_prompt("voice-test")
        assert "MODE VOCAL" in prompt_voice, "Suffixe voix non injecté"
        assert "MODE VOCAL" not in prompt_default, "Suffixe voix injecté à tort"
        return f"{len(TOOLS)} schemas, voice prompt OK"
    r.test("import server.orchestrator (TOOLS, voice prompt)", import_orchestrator)

    def schemas_match_handlers():
        from server.tools import ALL_HANDLERS
        from server.orchestrator import TOOLS
        handlers = set(ALL_HANDLERS.keys())
        schemas = set(t["name"] for t in TOOLS)
        # list_connected_devices est géré directement dans l'orchestrator
        schema_only = schemas - handlers - {"list_connected_devices"}
        handler_only = handlers - schemas
        assert not schema_only, f"Schema sans handler : {schema_only}"
        assert not handler_only, f"Handler sans schema : {handler_only}"
        return f"{len(handlers)} handlers ↔ {len(schemas)} schemas"
    r.test("schemas alignés avec handlers", schemas_match_handlers)


def test_filesystem(r: Runner):
    r.section("Tools filesystem")
    from server.tools import ALL_HANDLERS

    tmp = Path(tempfile.mkdtemp(prefix="orion_smoke_"))
    test_file = tmp / "hello.txt"
    test_dir = tmp / "subdir"
    moved_file = tmp / "renamed.txt"

    def create():
        out = ALL_HANDLERS["create_file"]({"path": str(test_file), "content": "hello orion"})
        assert out.get("success"), out
        assert test_file.exists()
    r.test("create_file", create)

    def read():
        out = ALL_HANDLERS["read_file"]({"path": str(test_file)})
        assert out.get("success"), out
        assert "hello orion" in out.get("content", "")
    r.test("read_file", read)

    def mkdir():
        out = ALL_HANDLERS["create_directory"]({"path": str(test_dir)})
        assert out.get("success"), out
        assert test_dir.is_dir()
    r.test("create_directory", mkdir)

    def list_dir():
        out = ALL_HANDLERS["list_directory"]({"path": str(tmp)})
        assert out.get("success"), out
        items = out.get("items", out.get("contents", []))
        names = [(i.get("name") if isinstance(i, dict) else i) for i in items]
        assert "hello.txt" in str(names), f"hello.txt pas dans {names}"
    r.test("list_directory", list_dir)

    def move():
        out = ALL_HANDLERS["move_file"]({"src": str(test_file), "dst": str(moved_file)})
        assert out.get("success"), out
        assert moved_file.exists() and not test_file.exists()
    r.test("move_file", move)

    def delete():
        out = ALL_HANDLERS["delete_file"]({"path": str(moved_file)})
        assert out.get("success"), out
        assert not moved_file.exists()
        # Cleanup
        ALL_HANDLERS["delete_file"]({"path": str(test_dir)})
        try:
            tmp.rmdir()
        except OSError:
            pass
    r.test("delete_file", delete)

    def sysinfo():
        out = ALL_HANDLERS["get_system_info"]({})
        assert out.get("success"), out
        info = out.get("info") or out
        # Vérifie qu'on a au moins l'OS et python_version dans la réponse
        flat = json.dumps(info)
        assert "Windows" in flat or "Linux" in flat or "Darwin" in flat, flat[:200]
    r.test("get_system_info", sysinfo)


def test_shell(r: Runner):
    r.section("Tools shell & code")
    from server.tools import ALL_HANDLERS

    def shell_echo():
        out = ALL_HANDLERS["run_shell_command"]({"command": "echo orion-smoke"})
        assert out.get("success"), out
        # output peut être dans 'stdout' ou 'output'
        text = out.get("stdout", "") + out.get("output", "")
        assert "orion-smoke" in text, f"output: {text[:200]}"
    r.test("run_shell_command (echo)", shell_echo)

    def python_script():
        out = ALL_HANDLERS["run_python_script"]({"code": "print(2 + 40)"})
        assert out.get("success"), out
        text = out.get("stdout", "") + out.get("output", "")
        assert "42" in text, f"output: {text[:200]}"
    r.test("run_python_script", python_script)


def test_web(r: Runner, online: bool):
    r.section("Tools web")
    from server.tools import ALL_HANDLERS

    if not online:
        for name in ("web_search", "fetch_url"):
            r.test(name, lambda: (_ for _ in ()).throw(SkipTest("--no-network")))
        return

    def search():
        out = ALL_HANDLERS["web_search"]({"query": "python language", "max_results": 3})
        assert out.get("success"), out
        results = out.get("results", [])
        assert len(results) >= 1, f"Aucun résultat : {out}"
        return f"{len(results)} résultats"
    r.test("web_search", search)

    def fetch():
        out = ALL_HANDLERS["fetch_url"]({"url": "https://example.com", "max_chars": 1000})
        assert out.get("success"), out
        content = out.get("content", "") + out.get("text", "")
        assert "Example" in content, content[:200]
    r.test("fetch_url (example.com)", fetch)


def test_notifications(r: Runner):
    r.section("Notifications")
    from server.tools import ALL_HANDLERS

    def notify():
        # winotify n'est pas forcément installé, on accepte les deux issues
        out = ALL_HANDLERS["notify"]({"title": "Orion smoke test", "message": "Test"})
        if not out.get("success") and "winotify" in out.get("error", ""):
            raise SkipTest("winotify non installé (pip install -r requirements-extras.txt)")
        assert out.get("success"), out
    r.test("notify (toast)", notify)


def test_screenshot(r: Runner):
    r.section("Capture d'écran")
    from server.tools import ALL_HANDLERS

    def list_mons():
        out = ALL_HANDLERS["list_monitors"]({})
        if not out.get("success") and "mss" in out.get("error", ""):
            raise SkipTest("mss non installé (pip install -r requirements-extras.txt)")
        assert out.get("success"), out
        assert len(out.get("monitors", [])) >= 1
        return f"{len(out['monitors'])} écran(s)"
    r.test("list_monitors", list_mons)

    def shot():
        with tempfile.TemporaryDirectory() as td:
            target = Path(td) / "shot.png"
            out = ALL_HANDLERS["screenshot"]({"path": str(target)})
            if not out.get("success") and "mss" in out.get("error", ""):
                raise SkipTest("mss non installé")
            assert out.get("success"), out
            assert target.exists() and target.stat().st_size > 1000, "PNG vide ou minuscule"
            return f"{target.stat().st_size // 1024} KB"
    r.test("screenshot (full screen)", shot)


def test_documents(r: Runner):
    r.section("Documents (PDF + DOCX)")
    from server.tools import ALL_HANDLERS

    # On crée un .docx temporaire pour tester read_docx
    def docx_roundtrip():
        try:
            from docx import Document
        except ImportError:
            raise SkipTest("python-docx non installé")
        with tempfile.TemporaryDirectory() as td:
            path = Path(td) / "test.docx"
            doc = Document()
            doc.add_paragraph("Ligne un de test Orion.")
            doc.add_paragraph("Ligne deux avec accents éàù.")
            doc.save(str(path))
            out = ALL_HANDLERS["read_docx"]({"path": str(path)})
            assert out.get("success"), out
            text = out.get("text", "")
            assert "Ligne un" in text and "éàù" in text, text[:200]
            return f"{len(text)} chars extraits"
    r.test("read_docx (round-trip)", docx_roundtrip)

    def pdf_lookup():
        # Pas de pdf de test fourni : on cherche dans le projet ou skip
        candidates = list(ROOT.glob("*.pdf")) + list((ROOT / "data").glob("**/*.pdf"))
        if not candidates:
            raise SkipTest("aucun PDF de test sous le projet")
        path = candidates[0]
        out = ALL_HANDLERS["read_pdf"]({"path": str(path), "max_chars": 500})
        if not out.get("success") and "pypdf" in out.get("error", ""):
            raise SkipTest("pypdf non installé")
        assert out.get("success"), out
        return f"{path.name} ({out.get('total_pages', '?')} pages)"
    r.test("read_pdf (premier PDF trouvé)", pdf_lookup)


def test_memory(r: Runner):
    r.section("Mémoire RAG")
    try:
        from server.memory.rag_tools import HANDLERS as M
    except ImportError as exc:
        for name in ("memory_remember", "memory_recall", "memory_stats", "memory_clear"):
            r.test(name, lambda: (_ for _ in ()).throw(SkipTest(f"import: {exc}")))
        return

    NS = "smoke_test_ns"

    # On vide d'abord pour partir propre
    try:
        M["memory_clear"]({"namespace": NS, "confirm": True})
    except Exception:
        pass

    item_id_holder = {}

    def remember():
        try:
            out = M["memory_remember"]({
                "text": "Le café noir sans sucre est apprécié de l'utilisateur.",
                "source": "smoke",
                "namespace": NS,
            })
        except ImportError:
            raise SkipTest("sentence-transformers non installé (pip install -r requirements-rag.txt)")
        assert out.get("success"), out
        item_id_holder["id"] = out.get("id")
        return out.get("id")
    r.test("memory_remember", remember)

    def recall():
        if "id" not in item_id_holder:
            raise SkipTest("dépend de memory_remember")
        out = M["memory_recall"]({"query": "préférences café", "top_k": 3, "namespace": NS})
        assert out.get("success"), out
        results = out.get("results", [])
        assert len(results) >= 1, f"Aucun résultat : {out}"
        assert "café" in results[0]["text"].lower() or "cafe" in results[0]["text"].lower()
        return f"top score = {results[0]['score']}"
    r.test("memory_recall (sémantique)", recall)

    def stats():
        out = M["memory_stats"]({"namespace": NS})
        assert out.get("success"), out
        assert out.get("count", 0) >= 1
        return f"count={out['count']}"
    r.test("memory_stats", stats)

    def index_file():
        with tempfile.NamedTemporaryFile(
            "w", suffix=".txt", delete=False, encoding="utf-8"
        ) as f:
            f.write("Orion est l'IA personnelle de l'utilisateur.\n" * 20)
            tmp_path = f.name
        try:
            out = M["memory_index_file"]({"path": tmp_path, "namespace": NS})
            assert out.get("success"), out
            assert out.get("chunks_added", 0) >= 1
            return f"{out['chunks_added']} chunks"
        finally:
            os.unlink(tmp_path)
    r.test("memory_index_file (.txt)", index_file)

    def forget():
        if "id" not in item_id_holder:
            raise SkipTest("dépend de memory_remember")
        out = M["memory_forget"]({"item_id": item_id_holder["id"], "namespace": NS})
        assert out.get("success"), out
    r.test("memory_forget", forget)

    def clear_ns():
        out = M["memory_clear"]({"namespace": NS, "confirm": True})
        assert out.get("success"), out
        return f"deleted {out.get('deleted', 0)}"
    r.test("memory_clear (cleanup)", clear_ns)


def test_image_gen(r: Runner, enabled: bool):
    r.section("Génération d'images")
    if not enabled:
        r.test("generate_image", lambda: (_ for _ in ()).throw(
            SkipTest("désactivé par défaut (passe --image-gen pour l'inclure, coûte 1 appel Gemini)")
        ))
        return
    from server.tools import ALL_HANDLERS

    def gen():
        if not os.getenv("GEMINI_API_KEY"):
            raise SkipTest("GEMINI_API_KEY manquant")
        out = ALL_HANDLERS["generate_image"]({
            "prompt": "Un petit logo géométrique minimaliste sur fond blanc",
            "n": 1,
            "aspect_ratio": "1:1",
        })
        assert out.get("success"), out
        paths = out.get("paths", [])
        assert paths and Path(paths[0]).exists()
        size = Path(paths[0]).stat().st_size
        return f"{Path(paths[0]).name} ({size // 1024} KB)"
    r.test("generate_image (Gemini Imagen)", gen)


def test_presets(r: Runner):
    r.section("Presets")
    presets_dir = ROOT / "presets"
    if not presets_dir.exists():
        r.test("presets dir", lambda: (_ for _ in ()).throw(SkipTest("presets/ absent")))
        return

    try:
        import tomllib  # type: ignore[import-not-found]
    except ImportError:
        try:
            import tomli as tomllib  # type: ignore[import-not-found]
        except ImportError:
            r.test("presets parse", lambda: (_ for _ in ()).throw(SkipTest("tomllib indisponible")))
            return

    for path in sorted(presets_dir.glob("*.toml")):
        def make_test(p):
            def t():
                with p.open("rb") as f:
                    data = tomllib.load(f)
                assert "name" in data, "champ 'name' manquant"
                assert isinstance(data.get("env", {}), dict)
                return data.get("description", "")[:40]
            return t
        r.test(f"preset: {path.stem}.toml", make_test(path))


# ─── Tests serveur (HTTP) ──────────────────────────────────────────────
def server_url() -> str:
    port = os.getenv("SERVER_PORT", "8765")
    return f"http://127.0.0.1:{port}"


def server_token() -> str:
    return get_env("SECRET_TOKEN", "") or ""


def is_server_up() -> bool:
    try:
        with urllib.request.urlopen(server_url() + "/status", timeout=1.5) as resp:
            return resp.status == 200
    except (urllib.error.URLError, OSError):
        return False


def test_server_endpoints(r: Runner):
    r.section("Endpoints serveur HTTP")
    if not is_server_up():
        for name in ("GET /status", "GET /devices", "POST /api/transcribe"):
            r.test(name, lambda: (_ for _ in ()).throw(
                SkipTest("serveur non joignable — lance `python start.py server` d'abord")
            ))
        return

    def status():
        with urllib.request.urlopen(server_url() + "/status", timeout=2) as resp:
            data = json.loads(resp.read())
        assert "online" in data.get("status", "").lower(), data
        return data.get("status")
    r.test("GET /status", status)

    def devices():
        token = server_token()
        if not token:
            raise SkipTest("ORION_SECRET_TOKEN absent")
        req = urllib.request.Request(
            server_url() + "/devices",
            headers={"Authorization": f"Bearer {token}"},
        )
        with urllib.request.urlopen(req, timeout=2) as resp:
            data = json.loads(resp.read())
        assert "controllers" in data and "workers" in data
        return f"controllers={len(data['controllers'])}, workers={len(data['workers'])}"
    r.test("GET /devices (avec token)", devices)

    def transcribe():
        token = server_token()
        if not token:
            raise SkipTest("ORION_SECRET_TOKEN absent")
        # Génère un WAV de silence valide (44 bytes header + 1s de silence à 16kHz mono int16)
        wav = _make_silent_wav(seconds=1, sample_rate=16000)
        url = f"{server_url()}/api/transcribe?token={token}&language=fr"
        req = urllib.request.Request(
            url, data=wav, method="POST",
            headers={"Content-Type": "audio/wav"},
        )
        try:
            with urllib.request.urlopen(req, timeout=30) as resp:
                data = json.loads(resp.read())
        except urllib.error.HTTPError as e:
            body = e.read().decode("utf-8", errors="replace")
            raise AssertionError(f"HTTP {e.code} : {body[:200]}")
        # Si faster-whisper pas installé, on a une 503 — capturé ci-dessus
        # Ici on attend un success avec text vide ou très court (silence)
        assert data.get("success"), data
        return f"text={data.get('text', '')[:30]!r}"
    r.test("POST /api/transcribe (silence WAV)", transcribe)


def _make_silent_wav(seconds: int = 1, sample_rate: int = 16000) -> bytes:
    """Génère un WAV mono 16-bit silencieux."""
    import struct
    n_samples = seconds * sample_rate
    data_size = n_samples * 2  # 2 bytes per int16 sample
    riff_size = 36 + data_size
    header = b"RIFF" + struct.pack("<I", riff_size) + b"WAVE"
    fmt = b"fmt " + struct.pack(
        "<IHHIIHH",
        16,             # subchunk1 size
        1,              # PCM
        1,              # mono
        sample_rate,
        sample_rate * 2,  # byte rate
        2,              # block align
        16,             # bits per sample
    )
    data = b"data" + struct.pack("<I", data_size) + b"\x00\x00" * n_samples
    return header + fmt + data


# ═══════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--no-network", action="store_true", help="skip web et endpoints serveur")
    parser.add_argument("--image-gen", action="store_true", help="inclut generate_image (1 appel Gemini)")
    parser.add_argument("--server-only", action="store_true", help="uniquement les tests d'endpoints serveur")
    args = parser.parse_args()

    print(f"{CYAN}╔══════════════════════════════════════════╗")
    print(f"║   ORION — Tests de fumée non-vocaux      ║")
    print(f"╚══════════════════════════════════════════╝{RESET}")

    r = Runner()

    if args.server_only:
        test_server_endpoints(r)
    else:
        test_imports_and_registry(r)
        test_filesystem(r)
        test_shell(r)
        test_web(r, online=not args.no_network)
        test_notifications(r)
        test_screenshot(r)
        test_documents(r)
        test_memory(r)
        test_image_gen(r, enabled=args.image_gen)
        test_presets(r)
        if not args.no_network:
            test_server_endpoints(r)

    return r.report()


if __name__ == "__main__":
    sys.exit(main())
