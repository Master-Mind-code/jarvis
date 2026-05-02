"""
Orion Tool — Lecture de documents (PDF + Word).

PDF  : pypdf (texte natif uniquement, pas d'OCR pour les scans)
DOCX : python-docx (paragraphes + tableaux)
"""
from __future__ import annotations

from pathlib import Path


def read_pdf(path: str, max_chars: int = 8000, pages: str | None = None) -> dict:
    """Extrait le texte d'un PDF. pages = '1-5' ou '3' ou None pour tout."""
    p = Path(path).expanduser()
    if not p.exists():
        return {"success": False, "error": f"Fichier introuvable : {p}"}

    try:
        from pypdf import PdfReader
    except ImportError:
        return {
            "success": False,
            "error": "pypdf n'est pas installé. Installe avec :\n"
                     "    pip install -r requirements-extras.txt",
        }

    try:
        reader = PdfReader(str(p))
    except Exception as exc:
        return {"success": False, "error": f"Lecture PDF impossible : {exc}"}

    n_pages = len(reader.pages)

    # Range de pages à lire
    page_indices = list(range(n_pages))
    if pages:
        try:
            if "-" in pages:
                a, b = pages.split("-", 1)
                a_i, b_i = max(1, int(a)), min(n_pages, int(b))
                page_indices = list(range(a_i - 1, b_i))
            else:
                idx = int(pages) - 1
                if 0 <= idx < n_pages:
                    page_indices = [idx]
        except ValueError:
            return {"success": False, "error": f"Format pages invalide : {pages!r}"}

    chunks: list[str] = []
    for i in page_indices:
        try:
            text = reader.pages[i].extract_text() or ""
        except Exception:
            text = ""
        if text.strip():
            chunks.append(f"--- Page {i + 1} ---\n{text.strip()}")

    full = "\n\n".join(chunks)
    truncated = False
    if len(full) > max_chars:
        full = full[:max_chars] + "\n[…tronqué]"
        truncated = True

    return {
        "success": True,
        "path": str(p),
        "total_pages": n_pages,
        "pages_read": len(page_indices),
        "text": full,
        "truncated": truncated,
        "metadata": {
            "title": (reader.metadata.title if reader.metadata else None) or "",
            "author": (reader.metadata.author if reader.metadata else None) or "",
        },
    }


def read_docx(path: str, max_chars: int = 8000) -> dict:
    """Extrait le texte d'un fichier Word (.docx). Inclut paragraphes + tableaux."""
    p = Path(path).expanduser()
    if not p.exists():
        return {"success": False, "error": f"Fichier introuvable : {p}"}

    try:
        from docx import Document  # type: ignore[import-not-found]
    except ImportError:
        return {
            "success": False,
            "error": "python-docx n'est pas installé. Installe avec :\n"
                     "    pip install -r requirements-extras.txt",
        }

    try:
        doc = Document(str(p))
    except Exception as exc:
        return {"success": False, "error": f"Lecture .docx impossible : {exc}"}

    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for i, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            parts.append(f"\n[Tableau {i + 1}]\n" + "\n".join(rows))

    full = "\n\n".join(parts)
    truncated = False
    if len(full) > max_chars:
        full = full[:max_chars] + "\n[…tronqué]"
        truncated = True

    return {
        "success": True,
        "path": str(p),
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "text": full,
        "truncated": truncated,
    }


HANDLERS = {
    "read_pdf": lambda p: read_pdf(**p),
    "read_docx": lambda p: read_docx(**p),
}
